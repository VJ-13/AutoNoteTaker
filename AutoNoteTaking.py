import requests
import sys
from time import sleep
import os.path
from docx import Document
from time import gmtime, strftime, localtime
import tkinter as tk
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
import ctypes


#input dialog box
from tkinter import simpledialog
ROOT = tk.Tk()
ROOT.withdraw()
fileID = 'RecordingOne'
fileID = simpledialog.askstring(title="Hi", prompt="Please create a file name for the transcript (e.g. RecordingOne):")

path=''
path = simpledialog.askstring(title="Hi", prompt="Please enter the path to the file that you want to transcript(mp3 or mp4 format):")


file_exists = os.path.exists(path)
if file_exists == 0:
    sys.exit("[ Please Cheack Your Path and Try Again ]")

def Mbox(title, text, style):
    return ctypes.windll.user32.MessageBoxW(0, text, title, style)
Mbox('All Set!', 'Your transcript will be ready in minutes', 1)

# store global constants
headers = {
   "authorization": "API key",
   "content-type": "application/json"
}
transcript_endpoint = "https://api.assemblyai.com/v2/transcript"
upload_endpoint = 'https://api.assemblyai.com/v2/upload'
 
# make a function to pass the mp3 to the upload endpoint
def read_file(filename):
   with open(filename, 'rb') as _file:
       while True:
           data = _file.read(5242880)
           if not data:
               break
           yield data
 
# upload our audio file
upload_response = requests.post(
   upload_endpoint,
   headers=headers, data=read_file(path)
)
print('Audio file uploaded')


# send a request to transcribe the audio file
transcript_request = {'audio_url': upload_response.json()['upload_url'], 'iab_categories': 'True', 'auto_highlights': True, "entity_detection": True, "auto_chapters": True}
transcript_response = requests.post(transcript_endpoint, json=transcript_request, headers=headers)
print('Transcription Requested, please wait...')

# set up polling
polling_response = requests.get(transcript_endpoint+"/"+transcript_response.json()['id'], headers=headers)
filename = fileID

# if our status isn’t complete, sleep and then poll again
while polling_response.json()['status'] != 'completed':
   sleep(20)
   polling_response = requests.get(transcript_endpoint+"/"+transcript_response.json()['id'], headers=headers)
   print("Transcript is", polling_response.json()['status'])

#get topics
topic = polling_response.json()["iab_categories_result"]["results"][0]["labels"][0]["label"]

#get highlights
hl=polling_response.json()["auto_highlights_result"]["results"][0]["text"]

#print entities
for i in range(len(polling_response.json()["entities"])):
   print(polling_response.json()["entities"][i]["entity_type"] + " : " + polling_response.json()["entities"][i]["text"])
   
#word document output
rawtext = polling_response.json()['text']
rawtext = rawtext.replace(". ", ".\n")

now = strftime("%Y-%m-%d %H:%M:%S", localtime())
wordoc = Document()
docOBJ = wordoc.add_paragraph(now+"\n")
docOBJ.add_run('\n'+'TOPICS:'+'\n')
docOBJ.add_run(topic+'\n'+'Hightlights:')
docOBJ.add_run(hl).font.highlight_color = WD_COLOR_INDEX.YELLOW
docOBJ.add_run('\n\n'+ 'Entities:'+'\n')

#insert entities
for i in range(len(polling_response.json()["entities"])):
   docOBJ.add_run(polling_response.json()["entities"][i]["entity_type"] + " : " + polling_response.json()["entities"][i]["text"]+'\n')
   
docOBJ.add_run('\n\n' + '=================================================================' + '\n\n')
docOBJ.add_run(rawtext)

docOBJ.add_run('\n\n'+'Auto Summary'+'\n').bold = True
for i in range(len(polling_response.json()["chapters"])):
   docOBJ.add_run(polling_response.json()["chapters"][i]["summary"]).font.color.rgb = RGBColor(0, 0, 204)
   
wordoc.save(filename + '.docx')




